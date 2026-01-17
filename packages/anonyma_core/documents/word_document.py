"""
Word document handler (.docx format).

Handles Word documents with:
- Text extraction from paragraphs, tables, headers, and footers
- Anonymization of content
- Document reconstruction preserving formatting
- Support for .docx format (OpenXML)
"""

from pathlib import Path
from typing import Dict, Any, List, Optional
from datetime import datetime
import io

from .base import BaseDocument, DocumentFormat, DocumentMetadata
from ..logging_config import get_logger
from ..exceptions import DocumentProcessingError

logger = get_logger(__name__)


class WordDocument(BaseDocument):
    """
    Word document handler for .docx files.

    Supports:
    - Text extraction from paragraphs
    - Text extraction from tables
    - Text extraction from headers/footers
    - Document reconstruction with anonymized content
    - Metadata preservation

    Note: Requires python-docx library.
    """

    def __init__(self, file_path: Path):
        """
        Initialize Word document handler.

        Args:
            file_path: Path to Word file

        Raises:
            DocumentProcessingError: If Word document can't be loaded
        """
        self._doc = None
        self._paragraphs_map = []  # Track original paragraph structure

        super().__init__(file_path)

        # Load document
        self._load_document()

    def _load_document(self):
        """Load Word document with python-docx"""
        try:
            from docx import Document

            self._doc = Document(self.file_path)

            logger.info(
                f"Word document loaded successfully",
                extra={
                    "extra_fields": {
                        "paragraphs": len(self._doc.paragraphs),
                        "tables": len(self._doc.tables),
                        "file_size": self.get_file_size(),
                    }
                },
            )

        except ImportError:
            raise DocumentProcessingError(
                "python-docx not installed. Install with: pip install python-docx",
                {"required_package": "python-docx"},
            )
        except Exception as e:
            logger.error(f"Failed to load Word document: {e}", exc_info=True)
            raise DocumentProcessingError(
                f"Failed to load Word document: {str(e)}",
                {"file_path": str(self.file_path)}
            )

    def _detect_format(self) -> DocumentFormat:
        """Detect format (always WORD for this handler)"""
        return DocumentFormat.WORD

    def _extract_metadata(self) -> DocumentMetadata:
        """Extract Word document metadata"""
        file_stats = self.file_path.stat()

        # Extract document properties
        author = None
        title = None
        creation_date = None
        modification_date = None

        try:
            if hasattr(self._doc, "core_properties"):
                props = self._doc.core_properties
                author = props.author
                title = props.title
                creation_date = props.created
                modification_date = props.modified

        except Exception as e:
            logger.warning(f"Failed to extract Word metadata: {e}")

        # Count paragraphs
        paragraph_count = len(self._doc.paragraphs) if self._doc else 0
        table_count = len(self._doc.tables) if self._doc else 0

        return DocumentMetadata(
            file_name=self.file_path.name,
            file_size=file_stats.st_size,
            format=DocumentFormat.WORD,
            author=author,
            title=title,
            creation_date=creation_date,
            modification_date=modification_date,
            is_scanned=False,  # .docx files are digital
            custom={
                "paragraphs": paragraph_count,
                "tables": table_count,
            }
        )

    def extract_text(self) -> str:
        """
        Extract all text from Word document.

        Extracts from:
        - Paragraphs
        - Tables
        - Headers and footers

        Returns:
            Extracted text content

        Raises:
            DocumentProcessingError: If extraction fails
        """
        try:
            text_parts = []

            # Extract from paragraphs
            logger.debug("Extracting text from paragraphs")
            for para in self._doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
                    self._paragraphs_map.append({
                        "type": "paragraph",
                        "text": para.text,
                        "index": len(text_parts) - 1
                    })

            # Extract from tables
            logger.debug(f"Extracting text from {len(self._doc.tables)} tables")
            for table_idx, table in enumerate(self._doc.tables):
                table_texts = []
                for row in table.rows:
                    row_texts = [cell.text.strip() for cell in row.cells]
                    if any(row_texts):
                        table_texts.append(" | ".join(row_texts))

                if table_texts:
                    table_text = "\n".join(table_texts)
                    text_parts.append(f"[Table {table_idx + 1}]\n{table_text}")
                    self._paragraphs_map.append({
                        "type": "table",
                        "text": table_text,
                        "index": len(text_parts) - 1,
                        "table_idx": table_idx
                    })

            # Extract from headers/footers
            logger.debug("Extracting text from headers and footers")
            for section in self._doc.sections:
                # Header
                if section.header:
                    for para in section.header.paragraphs:
                        if para.text.strip():
                            text_parts.append(f"[Header] {para.text}")

                # Footer
                if section.footer:
                    for para in section.footer.paragraphs:
                        if para.text.strip():
                            text_parts.append(f"[Footer] {para.text}")

            full_text = "\n\n".join(text_parts)

            logger.info(
                f"Text extraction completed",
                extra={
                    "extra_fields": {
                        "total_paragraphs": len(self._doc.paragraphs),
                        "total_tables": len(self._doc.tables),
                        "total_length": len(full_text),
                    }
                },
            )

            return full_text

        except Exception as e:
            logger.error(f"Text extraction failed: {e}", exc_info=True)
            raise DocumentProcessingError(
                f"Failed to extract text from Word document: {str(e)}",
                {"file_path": str(self.file_path)}
            )

    def rebuild(self, anonymized_text: str, detections: List[Dict[str, Any]]) -> bytes:
        """
        Rebuild Word document with anonymized content.

        Creates a new document with anonymized text, attempting to preserve
        basic structure (paragraphs, tables).

        Args:
            anonymized_text: Anonymized text
            detections: List of detections (for reference)

        Returns:
            Word document bytes
        """
        try:
            from docx import Document
            from docx.shared import Pt

            logger.info("Rebuilding Word document with anonymized content")

            # Create new document
            new_doc = Document()

            # Copy styles from original if possible
            try:
                if hasattr(self._doc, "styles"):
                    # Basic style copying - can be enhanced
                    pass
            except Exception as e:
                logger.debug(f"Could not copy styles: {e}")

            # Split anonymized text back into sections
            sections = anonymized_text.split("\n\n")

            for section in sections:
                if not section.strip():
                    continue

                # Check if it's a table section
                if section.startswith("[Table"):
                    # Extract table content
                    lines = section.split("\n")[1:]  # Skip [Table N] line
                    if lines:
                        # Create table
                        rows = [line.split(" | ") for line in lines if line.strip()]
                        if rows:
                            max_cols = max(len(row) for row in rows)
                            table = new_doc.add_table(rows=len(rows), cols=max_cols)
                            table.style = 'Table Grid'

                            for i, row_data in enumerate(rows):
                                for j, cell_text in enumerate(row_data):
                                    table.rows[i].cells[j].text = cell_text

                elif section.startswith("[Header]"):
                    # Add as paragraph with note
                    para = new_doc.add_paragraph(section)
                    para.runs[0].font.size = Pt(10)
                    para.runs[0].italic = True

                elif section.startswith("[Footer]"):
                    # Add as paragraph with note
                    para = new_doc.add_paragraph(section)
                    para.runs[0].font.size = Pt(10)
                    para.runs[0].italic = True

                else:
                    # Regular paragraph
                    new_doc.add_paragraph(section)

            # Save to bytes
            buffer = io.BytesIO()
            new_doc.save(buffer)
            doc_bytes = buffer.getvalue()
            buffer.close()

            logger.info(
                "Word document rebuilt successfully",
                extra={"extra_fields": {"output_size": len(doc_bytes)}},
            )

            return doc_bytes

        except Exception as e:
            logger.error(f"Word document rebuild failed: {e}", exc_info=True)
            raise DocumentProcessingError(
                f"Failed to rebuild Word document: {str(e)}",
                {"file_path": str(self.file_path)}
            )

    def get_paragraph_count(self) -> int:
        """
        Get number of paragraphs.

        Returns:
            Paragraph count
        """
        if hasattr(self, "_doc") and self._doc:
            return len(self._doc.paragraphs)
        return 0

    def get_table_count(self) -> int:
        """
        Get number of tables.

        Returns:
            Table count
        """
        if hasattr(self, "_doc") and self._doc:
            return len(self._doc.tables)
        return 0

    def close(self):
        """Close document resources"""
        if hasattr(self, "_doc"):
            self._doc = None
            logger.debug("Word document closed")

    def __del__(self):
        """Cleanup on deletion"""
        self.close()
