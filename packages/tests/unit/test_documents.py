"""
Unit tests for document processing module.

Tests for:
- BaseDocument and DocumentMetadata
- PDFDocument handler
- ImageDocument handler
- WordDocument handler
- ExcelDocument handler
- DocumentPipeline orchestration
- Custom pattern detector
"""

import pytest
from pathlib import Path
from datetime import datetime
import tempfile
import shutil

from anonyma_core.documents import (
    BaseDocument,
    DocumentFormat,
    DocumentMetadata,
    DocumentPipeline,
    ProcessingResult,
    PDFDocument,
    ImageDocument,
    WordDocument,
    ExcelDocument,
)
from anonyma_core.modes import AnonymizationMode
from anonyma_core.exceptions import DocumentProcessingError


# ============================================================================
# Fixtures
# ============================================================================


@pytest.fixture
def temp_dir():
    """Create temporary directory for test files"""
    tmpdir = tempfile.mkdtemp()
    yield Path(tmpdir)
    shutil.rmtree(tmpdir)


@pytest.fixture
def sample_text_italian():
    """Sample Italian text with PII"""
    return """
    Nome: Mario Rossi
    Email: mario.rossi@example.com
    Telefono: +39 339 1234567
    Codice Fiscale: RSSMRA85C15H501X
    """


@pytest.fixture
def mock_engine(mocker):
    """Mock AnonymaEngine for testing"""
    mock = mocker.MagicMock()

    # Mock anonymize method
    def mock_anonymize(text, mode, language):
        result = mocker.MagicMock()
        result.anonymized_text = "[REDACTED]"
        result.detections = [
            {"text": "Mario Rossi", "entity_type": "PERSON", "start": 10, "end": 21},
            {"text": "mario.rossi@example.com", "entity_type": "EMAIL", "start": 30, "end": 53},
        ]
        return result

    mock.anonymize = mock_anonymize
    return mock


# ============================================================================
# DocumentMetadata Tests
# ============================================================================


def test_document_metadata_creation():
    """Test DocumentMetadata creation"""
    metadata = DocumentMetadata(
        file_name="test.pdf",
        file_size=1024,
        format=DocumentFormat.PDF,
        page_count=5,
        author="Test Author",
    )

    assert metadata.file_name == "test.pdf"
    assert metadata.file_size == 1024
    assert metadata.format == DocumentFormat.PDF
    assert metadata.page_count == 5
    assert metadata.author == "Test Author"


def test_document_metadata_to_dict():
    """Test DocumentMetadata serialization"""
    metadata = DocumentMetadata(
        file_name="test.pdf",
        file_size=1024,
        format=DocumentFormat.PDF,
        is_scanned=True,
    )

    data = metadata.to_dict()

    assert data["file_name"] == "test.pdf"
    assert data["file_size"] == 1024
    assert data["format"] == "pdf"
    assert data["is_scanned"] is True


# ============================================================================
# DocumentFormat Tests
# ============================================================================


def test_document_formats():
    """Test DocumentFormat enum values"""
    assert DocumentFormat.PDF.value == "pdf"
    assert DocumentFormat.IMAGE.value == "image"
    assert DocumentFormat.WORD.value == "word"
    assert DocumentFormat.EXCEL.value == "excel"
    assert DocumentFormat.TEXT.value == "text"


# ============================================================================
# PDFDocument Tests
# ============================================================================


@pytest.mark.skipif(
    not pytest.importorskip("pdfplumber", reason="pdfplumber not installed"),
    reason="Requires pdfplumber",
)
def test_pdf_document_creation_requires_file():
    """Test PDFDocument requires valid file"""
    with pytest.raises(DocumentProcessingError):
        PDFDocument(Path("nonexistent.pdf"))


def test_pdf_document_detect_format():
    """Test PDF format detection"""
    # This would require a real PDF file
    # Skipping for now unless we create test fixtures
    pass


# ============================================================================
# ImageDocument Tests
# ============================================================================


@pytest.mark.skipif(
    not pytest.importorskip("PIL", reason="Pillow not installed"),
    reason="Requires Pillow",
)
def test_image_document_creation_requires_file():
    """Test ImageDocument requires valid file"""
    with pytest.raises(DocumentProcessingError):
        ImageDocument(Path("nonexistent.png"))


def test_image_document_format():
    """Test image format detection"""
    # Requires test image file
    pass


# ============================================================================
# WordDocument Tests
# ============================================================================


@pytest.mark.skipif(
    not pytest.importorskip("docx", reason="python-docx not installed"),
    reason="Requires python-docx",
)
def test_word_document_creation_requires_file():
    """Test WordDocument requires valid file"""
    with pytest.raises(DocumentProcessingError):
        WordDocument(Path("nonexistent.docx"))


@pytest.mark.skipif(
    not pytest.importorskip("docx", reason="python-docx not installed"),
    reason="Requires python-docx",
)
def test_word_document_creation(temp_dir):
    """Test WordDocument creation with actual file"""
    from docx import Document

    # Create sample Word document
    doc_path = temp_dir / "test.docx"
    doc = Document()
    doc.add_paragraph("Test paragraph with Mario Rossi")
    doc.save(str(doc_path))

    # Load with WordDocument handler
    word_doc = WordDocument(doc_path)

    assert word_doc.format == DocumentFormat.WORD
    assert word_doc.get_paragraph_count() > 0


@pytest.mark.skipif(
    not pytest.importorskip("docx", reason="python-docx not installed"),
    reason="Requires python-docx",
)
def test_word_document_extract_text(temp_dir):
    """Test Word text extraction"""
    from docx import Document

    # Create sample document
    doc_path = temp_dir / "test.docx"
    doc = Document()
    doc.add_paragraph("First paragraph")
    doc.add_paragraph("Second paragraph")
    doc.save(str(doc_path))

    # Extract text
    word_doc = WordDocument(doc_path)
    text = word_doc.extract_text()

    assert "First paragraph" in text
    assert "Second paragraph" in text


@pytest.mark.skipif(
    not pytest.importorskip("docx", reason="python-docx not installed"),
    reason="Requires python-docx",
)
def test_word_document_rebuild(temp_dir):
    """Test Word document rebuild"""
    from docx import Document

    # Create sample document
    doc_path = temp_dir / "test.docx"
    doc = Document()
    doc.add_paragraph("Test content")
    doc.save(str(doc_path))

    # Rebuild with anonymized content
    word_doc = WordDocument(doc_path)
    anonymized_bytes = word_doc.rebuild("[REDACTED]", [])

    assert isinstance(anonymized_bytes, bytes)
    assert len(anonymized_bytes) > 0


# ============================================================================
# ExcelDocument Tests
# ============================================================================


@pytest.mark.skipif(
    not pytest.importorskip("openpyxl", reason="openpyxl not installed"),
    reason="Requires openpyxl",
)
def test_excel_document_creation_requires_file():
    """Test ExcelDocument requires valid file"""
    with pytest.raises(DocumentProcessingError):
        ExcelDocument(Path("nonexistent.xlsx"))


@pytest.mark.skipif(
    not pytest.importorskip("openpyxl", reason="openpyxl not installed"),
    reason="Requires openpyxl",
)
def test_excel_document_creation(temp_dir):
    """Test ExcelDocument creation with actual file"""
    from openpyxl import Workbook

    # Create sample Excel document
    excel_path = temp_dir / "test.xlsx"
    wb = Workbook()
    ws = wb.active
    ws["A1"] = "Name"
    ws["B1"] = "Email"
    ws["A2"] = "Mario Rossi"
    ws["B2"] = "mario.rossi@example.com"
    wb.save(str(excel_path))

    # Load with ExcelDocument handler
    excel_doc = ExcelDocument(excel_path)

    assert excel_doc.format == DocumentFormat.EXCEL
    assert excel_doc.get_sheet_count() > 0


@pytest.mark.skipif(
    not pytest.importorskip("openpyxl", reason="openpyxl not installed"),
    reason="Requires openpyxl",
)
def test_excel_document_extract_text(temp_dir):
    """Test Excel text extraction"""
    from openpyxl import Workbook

    # Create sample document
    excel_path = temp_dir / "test.xlsx"
    wb = Workbook()
    ws = wb.active
    ws["A1"] = "Test"
    ws["B1"] = "Data"
    ws["A2"] = "Row1"
    ws["B2"] = "Value1"
    wb.save(str(excel_path))

    # Extract text
    excel_doc = ExcelDocument(excel_path)
    text = excel_doc.extract_text()

    assert "Test" in text
    assert "Data" in text
    assert "Row1" in text
    assert "Value1" in text


@pytest.mark.skipif(
    not pytest.importorskip("openpyxl", reason="openpyxl not installed"),
    reason="Requires openpyxl",
)
def test_excel_document_rebuild(temp_dir):
    """Test Excel document rebuild"""
    from openpyxl import Workbook

    # Create sample document
    excel_path = temp_dir / "test.xlsx"
    wb = Workbook()
    ws = wb.active
    ws["A1"] = "Test"
    wb.save(str(excel_path))

    # Rebuild with anonymized content
    excel_doc = ExcelDocument(excel_path)
    anonymized_bytes = excel_doc.rebuild("[Sheet: Sheet]\n[REDACTED]", [])

    assert isinstance(anonymized_bytes, bytes)
    assert len(anonymized_bytes) > 0


# ============================================================================
# DocumentPipeline Tests
# ============================================================================


def test_document_pipeline_initialization(mock_engine):
    """Test DocumentPipeline initialization"""
    pipeline = DocumentPipeline(mock_engine)

    assert pipeline.engine == mock_engine
    assert len(pipeline.handlers) > 0
    assert DocumentFormat.PDF in pipeline.handlers
    assert DocumentFormat.IMAGE in pipeline.handlers
    assert DocumentFormat.WORD in pipeline.handlers
    assert DocumentFormat.EXCEL in pipeline.handlers


def test_document_pipeline_supported_formats(mock_engine):
    """Test getting supported formats"""
    pipeline = DocumentPipeline(mock_engine)

    formats = pipeline.get_supported_formats()

    assert DocumentFormat.PDF in formats
    assert DocumentFormat.IMAGE in formats
    assert DocumentFormat.WORD in formats
    assert DocumentFormat.EXCEL in formats


def test_document_pipeline_is_format_supported(mock_engine):
    """Test format support checking"""
    pipeline = DocumentPipeline(mock_engine)

    assert pipeline.is_format_supported(DocumentFormat.PDF) is True
    assert pipeline.is_format_supported(DocumentFormat.IMAGE) is True
    assert pipeline.is_format_supported(DocumentFormat.WORD) is True
    assert pipeline.is_format_supported(DocumentFormat.EXCEL) is True
    assert pipeline.is_format_supported(DocumentFormat.UNKNOWN) is False


def test_document_pipeline_detect_format(mock_engine):
    """Test format detection"""
    pipeline = DocumentPipeline(mock_engine)

    # Test various extensions
    assert pipeline._detect_format(Path("test.pdf")) == DocumentFormat.PDF
    assert pipeline._detect_format(Path("test.png")) == DocumentFormat.IMAGE
    assert pipeline._detect_format(Path("test.jpg")) == DocumentFormat.IMAGE
    assert pipeline._detect_format(Path("test.docx")) == DocumentFormat.WORD
    assert pipeline._detect_format(Path("test.xlsx")) == DocumentFormat.EXCEL
    assert pipeline._detect_format(Path("test.txt")) == DocumentFormat.TEXT


def test_document_pipeline_unsupported_format(mock_engine):
    """Test handling unsupported format"""
    pipeline = DocumentPipeline(mock_engine)

    with pytest.raises(DocumentProcessingError):
        pipeline._get_handler(Path("test.xyz"), DocumentFormat.UNKNOWN)


@pytest.mark.skipif(
    not pytest.importorskip("docx", reason="python-docx not installed"),
    reason="Requires python-docx",
)
def test_document_pipeline_process_word(temp_dir, mock_engine):
    """Test full pipeline with Word document"""
    from docx import Document

    # Create sample Word document
    doc_path = temp_dir / "test.docx"
    doc = Document()
    doc.add_paragraph("Test with Mario Rossi")
    doc.save(str(doc_path))

    # Process through pipeline
    pipeline = DocumentPipeline(mock_engine)
    result = pipeline.process(
        file_path=doc_path,
        mode=AnonymizationMode.REDACT,
        output_path=temp_dir / "output.docx",
        language="it",
        save_output=True,
    )

    assert result.success is True
    assert result.format == DocumentFormat.WORD
    assert result.output_file is not None
    assert result.output_file.exists()


@pytest.mark.skipif(
    not pytest.importorskip("openpyxl", reason="openpyxl not installed"),
    reason="Requires openpyxl",
)
def test_document_pipeline_process_excel(temp_dir, mock_engine):
    """Test full pipeline with Excel document"""
    from openpyxl import Workbook

    # Create sample Excel document
    excel_path = temp_dir / "test.xlsx"
    wb = Workbook()
    ws = wb.active
    ws["A1"] = "Mario Rossi"
    wb.save(str(excel_path))

    # Process through pipeline
    pipeline = DocumentPipeline(mock_engine)
    result = pipeline.process(
        file_path=excel_path,
        mode=AnonymizationMode.REDACT,
        output_path=temp_dir / "output.xlsx",
        language="it",
        save_output=True,
    )

    assert result.success is True
    assert result.format == DocumentFormat.EXCEL
    assert result.output_file is not None
    assert result.output_file.exists()


# ============================================================================
# ProcessingResult Tests
# ============================================================================


def test_processing_result_creation():
    """Test ProcessingResult creation"""
    result = ProcessingResult(
        success=True,
        original_file=Path("test.pdf"),
        output_file=Path("output.pdf"),
        format=DocumentFormat.PDF,
        anonymized_text="[REDACTED]",
        detections_count=5,
        mode=AnonymizationMode.REDACT,
        processing_time=1.5,
    )

    assert result.success is True
    assert result.format == DocumentFormat.PDF
    assert result.detections_count == 5
    assert result.processing_time == 1.5


def test_processing_result_to_dict():
    """Test ProcessingResult serialization"""
    result = ProcessingResult(
        success=True,
        original_file=Path("test.pdf"),
        format=DocumentFormat.PDF,
        mode=AnonymizationMode.REDACT,
    )

    data = result.to_dict()

    assert data["success"] is True
    assert data["format"] == "pdf"
    assert data["mode"] == "redact"
    assert "original_file" in data


def test_processing_result_failure():
    """Test ProcessingResult for failed processing"""
    result = ProcessingResult(
        success=False,
        original_file=Path("test.pdf"),
        error="File not found",
    )

    assert result.success is False
    assert result.error == "File not found"
    assert result.output_file is None


# ============================================================================
# Integration Tests
# ============================================================================


@pytest.mark.skipif(
    not pytest.importorskip("docx", reason="python-docx not installed"),
    reason="Requires python-docx",
)
def test_full_workflow_word_document(temp_dir, mock_engine):
    """Test complete workflow for Word document"""
    from docx import Document

    # Create document with PII
    doc_path = temp_dir / "confidential.docx"
    doc = Document()
    doc.add_paragraph("Employee: Mario Rossi")
    doc.add_paragraph("Email: mario.rossi@example.com")
    doc.add_paragraph("Phone: +39 339 1234567")
    doc.save(str(doc_path))

    # Process
    pipeline = DocumentPipeline(mock_engine)
    result = pipeline.process(
        file_path=doc_path,
        mode=AnonymizationMode.REDACT,
        language="it",
    )

    assert result.success is True
    assert result.detections_count > 0
    assert result.anonymized_text is not None


@pytest.mark.skipif(
    not pytest.importorskip("openpyxl", reason="openpyxl not installed"),
    reason="Requires openpyxl",
)
def test_full_workflow_excel_document(temp_dir, mock_engine):
    """Test complete workflow for Excel document"""
    from openpyxl import Workbook

    # Create spreadsheet with PII
    excel_path = temp_dir / "contacts.xlsx"
    wb = Workbook()
    ws = wb.active
    ws["A1"] = "Name"
    ws["B1"] = "Email"
    ws["C1"] = "Phone"
    ws["A2"] = "Mario Rossi"
    ws["B2"] = "mario.rossi@example.com"
    ws["C2"] = "+39 339 1234567"
    wb.save(str(excel_path))

    # Process
    pipeline = DocumentPipeline(mock_engine)
    result = pipeline.process(
        file_path=excel_path,
        mode=AnonymizationMode.SUBSTITUTE,
        language="it",
    )

    assert result.success is True
    assert result.format == DocumentFormat.EXCEL
