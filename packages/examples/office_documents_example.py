#!/usr/bin/env python3
"""
Example: Office Documents Processing (Word & Excel)

Demonstrates how to:
1. Process Word documents (.docx)
2. Process Excel spreadsheets (.xlsx)
3. Anonymize PII in office documents
4. Preserve document structure
"""

import sys
from pathlib import Path

# Add parent directory to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from anonyma_core import AnonymaEngine
from anonyma_core.modes import AnonymizationMode
from anonyma_core.documents import DocumentPipeline


def create_sample_word_document():
    """Create a sample Word document for testing"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except ImportError:
        print("❌ python-docx not installed. Install with: pip install python-docx")
        return None

    # Sample content with PII
    output_path = Path("sample_hr_document.docx")

    try:
        doc = Document()

        # Title
        title = doc.add_heading("CONFIDENTIAL - Employee Information", level=1)

        # Personal Information Section
        doc.add_heading("Personal Information", level=2)
        doc.add_paragraph("Full Name: Mario Rossi")
        doc.add_paragraph("Date of Birth: 15/03/1985")
        doc.add_paragraph("Address: Via Roma 123, 20100 Milano (MI)")

        # Contact Section
        doc.add_heading("Contact Details", level=2)
        doc.add_paragraph("Email: mario.rossi@company.com")
        doc.add_paragraph("Phone: +39 339 1234567")
        doc.add_paragraph("Emergency Contact: +39 02 12345678")

        # Employment Section
        doc.add_heading("Employment Data", level=2)
        doc.add_paragraph("Employee ID: EMP-001234")
        doc.add_paragraph("Department: Research & Development")
        doc.add_paragraph("Hire Date: 01/01/2020")

        # Tax Information Section
        doc.add_heading("Tax Information", level=2)
        doc.add_paragraph("Fiscal Code: RSSMRA85C15H501X")
        doc.add_paragraph("VAT Number: 12345678901")

        # Table with sensitive data
        doc.add_heading("Salary History", level=2)
        table = doc.add_table(rows=4, cols=3)
        table.style = 'Table Grid'

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Year'
        header_cells[1].text = 'Position'
        header_cells[2].text = 'Salary'

        # Data rows
        data = [
            ('2020', 'Junior Developer', '€35,000'),
            ('2021', 'Developer', '€42,000'),
            ('2022', 'Senior Developer', '€52,000'),
        ]

        for i, (year, position, salary) in enumerate(data, start=1):
            row_cells = table.rows[i].cells
            row_cells[0].text = year
            row_cells[1].text = position
            row_cells[2].text = salary

        # Notes
        doc.add_heading("Notes", level=2)
        doc.add_paragraph(
            "Mr. Rossi has been an excellent employee. "
            "His email m.rossi@personal.it is on file for personal communication."
        )

        doc.save(str(output_path))
        print(f"✅ Sample Word document created: {output_path}")
        return output_path

    except Exception as e:
        print(f"❌ Failed to create Word document: {e}")
        return None


def create_sample_excel_spreadsheet():
    """Create a sample Excel spreadsheet for testing"""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill
    except ImportError:
        print("❌ openpyxl not installed. Install with: pip install openpyxl")
        return None

    output_path = Path("sample_customer_data.xlsx")

    try:
        wb = Workbook()

        # Sheet 1: Customer Data
        ws1 = wb.active
        ws1.title = "Customers"

        # Headers
        headers = ['Customer ID', 'Name', 'Email', 'Phone', 'Fiscal Code']
        for col, header in enumerate(headers, start=1):
            cell = ws1.cell(row=1, column=col, value=header)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")

        # Sample data
        customers = [
            ('CUST-001', 'Mario Rossi', 'mario.rossi@email.com', '+39 339 1234567', 'RSSMRA85C15H501X'),
            ('CUST-002', 'Laura Bianchi', 'laura.bianchi@email.com', '+39 340 2345678', 'BNCLAR90D45F205Y'),
            ('CUST-003', 'Giovanni Verdi', 'g.verdi@email.com', '+39 333 3456789', 'VRDGNN88E10H501Z'),
            ('CUST-004', 'Sara Neri', 'sara.neri@email.com', '+39 348 4567890', 'NRSSRA92M50F205W'),
        ]

        for row, customer in enumerate(customers, start=2):
            for col, value in enumerate(customer, start=1):
                ws1.cell(row=row, column=col, value=value)

        # Sheet 2: Orders
        ws2 = wb.create_sheet(title="Orders")

        # Headers
        order_headers = ['Order ID', 'Customer ID', 'Date', 'Amount', 'Status']
        for col, header in enumerate(order_headers, start=1):
            cell = ws2.cell(row=1, column=col, value=header)
            cell.font = Font(bold=True)

        # Sample orders
        orders = [
            ('ORD-10001', 'CUST-001', '2024-01-15', '€250.00', 'Completed'),
            ('ORD-10002', 'CUST-002', '2024-01-16', '€180.50', 'Completed'),
            ('ORD-10003', 'CUST-001', '2024-01-17', '€420.00', 'Processing'),
            ('ORD-10004', 'CUST-003', '2024-01-18', '€95.00', 'Shipped'),
        ]

        for row, order in enumerate(orders, start=2):
            for col, value in enumerate(order, start=1):
                ws2.cell(row=row, column=col, value=value)

        # Sheet 3: Notes
        ws3 = wb.create_sheet(title="Notes")
        ws3['A1'] = 'CONFIDENTIAL CUSTOMER NOTES'
        ws3['A1'].font = Font(bold=True, size=14)

        ws3['A3'] = 'Customer mario.rossi@email.com prefers email communication.'
        ws3['A4'] = 'Contact CUST-002 via phone only: +39 340 2345678'
        ws3['A5'] = 'VIP customer: Giovanni Verdi (VRDGNN88E10H501Z)'

        wb.save(str(output_path))
        print(f"✅ Sample Excel spreadsheet created: {output_path}")
        return output_path

    except Exception as e:
        print(f"❌ Failed to create Excel spreadsheet: {e}")
        return None


def process_word_example():
    """Example of processing a Word document"""
    print("=" * 70)
    print("WORD DOCUMENT PROCESSING")
    print("=" * 70)
    print()

    # Create sample document
    print("Step 1: Creating sample Word document...")
    print("-" * 70)
    word_path = create_sample_word_document()

    if not word_path:
        print("⚠️  Could not create sample Word document.")
        return

    print()

    # Initialize engine and pipeline
    print("Step 2: Initializing Anonyma...")
    print("-" * 70)

    try:
        engine = AnonymaEngine(use_flair=False)
        pipeline = DocumentPipeline(engine)
        print("✅ Engine initialized")
        print()

    except Exception as e:
        print(f"❌ Failed to initialize: {e}")
        return

    # Process with different modes
    print("Step 3: Processing Word document...")
    print("-" * 70)
    print()

    modes = [
        (AnonymizationMode.REDACT, "Redaction (blocks)"),
        (AnonymizationMode.SUBSTITUTE, "Substitution (fake data)"),
    ]

    for mode, description in modes:
        print(f"Processing with {description}...")

        try:
            result = pipeline.process(
                file_path=word_path,
                mode=mode,
                output_path=Path(f"anonymized_{mode.value}_{word_path.name}"),
                language="it",
            )

            if result.success:
                print(f"✅ Success!")
                print(f"   - Detections: {result.detections_count}")
                print(f"   - Time: {result.processing_time:.2f}s")
                print(f"   - Output: {result.output_file}")
                print(f"   - Paragraphs: {result.metadata.get('custom', {}).get('paragraphs', 'N/A')}")
            else:
                print(f"❌ Failed: {result.error}")

            print()

        except Exception as e:
            print(f"❌ Error: {e}")
            print()


def process_excel_example():
    """Example of processing an Excel spreadsheet"""
    print("=" * 70)
    print("EXCEL SPREADSHEET PROCESSING")
    print("=" * 70)
    print()

    # Create sample spreadsheet
    print("Step 1: Creating sample Excel spreadsheet...")
    print("-" * 70)
    excel_path = create_sample_excel_spreadsheet()

    if not excel_path:
        print("⚠️  Could not create sample Excel spreadsheet.")
        return

    print()

    # Initialize engine and pipeline
    print("Step 2: Initializing Anonyma...")
    print("-" * 70)

    try:
        engine = AnonymaEngine(use_flair=False)
        pipeline = DocumentPipeline(engine)
        print("✅ Engine initialized")
        print()

    except Exception as e:
        print(f"❌ Failed to initialize: {e}")
        return

    # Process
    print("Step 3: Processing Excel spreadsheet...")
    print("-" * 70)
    print()

    modes = [
        (AnonymizationMode.REDACT, "Redaction (blocks)"),
        (AnonymizationMode.SUBSTITUTE, "Substitution (fake data)"),
    ]

    for mode, description in modes:
        print(f"Processing with {description}...")

        try:
            result = pipeline.process(
                file_path=excel_path,
                mode=mode,
                output_path=Path(f"anonymized_{mode.value}_{excel_path.name}"),
                language="it",
            )

            if result.success:
                print(f"✅ Success!")
                print(f"   - Detections: {result.detections_count}")
                print(f"   - Time: {result.processing_time:.2f}s")
                print(f"   - Output: {result.output_file}")
                print(f"   - Sheets: {result.metadata.get('custom', {}).get('sheets', 'N/A')}")
                print(f"   - Cells: {result.metadata.get('custom', {}).get('total_cells', 'N/A')}")
            else:
                print(f"❌ Failed: {result.error}")

            print()

        except Exception as e:
            print(f"❌ Error: {e}")
            print()


def main():
    """Run all examples"""
    print()
    print("╔" + "=" * 68 + "╗")
    print("║" + " " * 15 + "OFFICE DOCUMENTS PROCESSING EXAMPLES" + " " * 17 + "║")
    print("╚" + "=" * 68 + "╝")
    print()
    print("Demonstrating anonymization of Word and Excel documents")
    print()

    try:
        # Word processing
        process_word_example()
        input("Press Enter to continue to Excel example...\n")

        # Excel processing
        process_excel_example()

        print("=" * 70)
        print("ALL EXAMPLES COMPLETED!")
        print("=" * 70)
        print()
        print("Generated files:")
        print("  - sample_hr_document.docx (original Word)")
        print("  - anonymized_*.docx (anonymized Word)")
        print("  - sample_customer_data.xlsx (original Excel)")
        print("  - anonymized_*.xlsx (anonymized Excel)")
        print()
        print("Key Features Demonstrated:")
        print("  ✓ Word document text extraction")
        print("  ✓ Excel multi-sheet processing")
        print("  ✓ Table handling in both formats")
        print("  ✓ Structure preservation")
        print("  ✓ Multiple anonymization modes")
        print()

    except KeyboardInterrupt:
        print("\n\n⚠️  Examples interrupted by user")
    except Exception as e:
        print(f"\n\n❌ Error: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
